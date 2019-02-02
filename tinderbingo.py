import random
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL

if __name__ == '__main__':

    document = Document()

    g = "N/A"

    while(g.lower() != "m" and g.lower() != "f" and g.lower() != "b"):
        g = input("Hvilket kjønn sveiper dere på? Mann, kvinne eller begge: (M/F/B): ")

    n = 0

    while(n < 1):
        n = int(input("Hvor mange spillere er dere? "))

    for k in range(n):

        board = [["" for x in range(4)] for y in range(4)]

        with open("bingo.txt") as f:
            bingolist = f.read().split('\n')

        used = []
        for i in range(4):
            for j in range(4):

                linenum = random.randint(1,104)
                while(bingolist[linenum-1] in used):
                    linenum = random.randint(1, 104)

                if g.lower() == "f":
                    while(69 < linenum < 88):
                        linenum = random.randint(1, 104)

                if g.lower() == "m":
                    while(linenum >= 88):
                        linenum = random.randint(1, 104)

                board[i][j] = bingolist[linenum-1]
                used.append(bingolist[linenum-1])

        #for i in range(4):
            # print('\n\n\n\n{:50s}{:50s}{:50s}{:50s}\n\n\n\n'.format(board[i][0],board[i][1],board[i][2],board[i][3]))

        document.add_heading('Tinderbingo! ' +  u'\U0001F525', 0)

        table = document.add_table(rows=0, cols=4)
        table.style = 'Table Grid'
        for row in board:
            row_cells = table.add_row().cells
            i = 0
            for col in row:
                row_cells[i].text = col
                i += 1

        for row in table.rows:
            row.height = Inches(1.5)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.width = Inches(1.5)

        if(n > 1 and k < n-1):
            document.add_page_break()

    print("Bingobrett generert i ./bingobrett.docx")
    document.save('bingobrett.docx')